from fastapi import APIRouter, HTTPException, Depends, Body, Query, UploadFile, File, Form
from typing import List, Dict, Any, Optional, Union
from datetime import datetime
from pydantic import BaseModel
from db import get_db_cursor
from admin_auth import require_admin
from user_auth import get_current_user
from openai_client import get_openai_client
from llm_gateway import chat_completions_create
import json
import os
import time
import io
import re
from functools import lru_cache
from urllib.request import Request, urlopen
from urllib.error import URLError, HTTPError

router = APIRouter(prefix="/api/contracts", tags=["Sözleşmeler"])

# --- Mock Data for Demo Templates ---
MOCK_TEMPLATES = [
    {
        "id": 1,
        "title": "İş Sözleşmesi (Belirsiz Süreli)",
        "category": "İş",
        "content": "Madde 1: Taraflar...\nMadde 2: Görev Tanımı...\nMadde 3: Ücret...\nMadde 4: Çalışma Süresi...",
        "description": "Standart personel iş sözleşmesi."
    },
    {
        "id": 2,
        "title": "Kira Sözleşmesi (Konut)",
        "category": "Kira",
        "content": "Madde 1: Kiralanan Yer...\nMadde 2: Kira Bedeli...\nMadde 3: Depozito...\nMadde 4: Kira Artışı...",
        "description": "TBK hükümlerine uygun konut kira kontratı."
    },
    {
        "id": 3,
        "title": "Gizlilik Sözleşmesi (NDA)",
        "category": "Ticari",
        "content": "Madde 1: Gizli Bilgi Tanımı...\nMadde 2: Tarafların Yükümlülükleri...\nMadde 3: Cezai Şart...",
        "description": "Şirketler arası bilgi paylaşımı için."
    },
     {
        "id": 4,
        "title": "Hizmet Alım Sözleşmesi",
        "category": "Ticari",
        "content": "Madde 1: Konu...\nMadde 2: Hizmetin Kapsamı...\nMadde 3: Ödeme Şartları...",
        "description": "Freelancer veya ajans hizmet alımları için."
    }
]

SEED_CATEGORIES = [
    {
        "key": "kira",
        "title": "Kira Sözleşmeleri",
        "icon": "🏠",
        "description": "Konut/işyeri kiralama, depozito ve tahliye protokolleri.",
    },
    {
        "key": "is",
        "title": "İş Sözleşmeleri",
        "icon": "👔",
        "description": "İşe giriş, çalışma koşulları, gizlilik, rekabet ve fesih süreçleri.",
    },
    {
        "key": "hizmet",
        "title": "Hizmet ve Danışmanlık",
        "icon": "🧩",
        "description": "Hizmet alımı, danışmanlık, bakım ve saha hizmetleri sözleşmeleri.",
    },
    {
        "key": "satis",
        "title": "Satış ve Tedarik",
        "icon": "🧾",
        "description": "Mal/hizmet satışı, tedarik, bayi/distribütör ve satış sonrası şartlar.",
    },
    {
        "key": "gizlilik",
        "title": "Gizlilik ve Fikri Haklar",
        "icon": "🔒",
        "description": "NDA, IP devri/lisans, marka ve içerik kullanım koşulları.",
    },
    {
        "key": "sirket",
        "title": "Şirket ve Ortaklık",
        "icon": "🏢",
        "description": "Hissedarlık/ortaklık, pay devri, yetki ve yönetim protokolleri.",
    },
    {
        "key": "teknoloji",
        "title": "Teknoloji ve Yazılım",
        "icon": "💻",
        "description": "SaaS, yazılım geliştirme, bakım, SLA ve veri işleme sözleşmeleri.",
    },
    {
        "key": "finans",
        "title": "Finans ve Teminat",
        "icon": "💳",
        "description": "Borç/ödünç, kefalet, teminat, ödeme planı ve tahsilat protokolleri.",
    },
    {
        "key": "lojistik",
        "title": "Lojistik ve Taşıma",
        "icon": "🚚",
        "description": "Taşıma, depolama, teslim ve hasar sorumluluğu düzenlemeleri.",
    },
    {
        "key": "egitim",
        "title": "Eğitim, Sağlık ve Etkinlik",
        "icon": "🎓",
        "description": "Kurs, seminer, klinik hizmet ve etkinlik organizasyonu sözleşmeleri.",
    },
    {
        "key": "pazarlama",
        "title": "Pazarlama ve Medya",
        "icon": "📣",
        "description": "Influencer, ajans, reklam yayın ve içerik üretim sözleşmeleri.",
    },
]

FIELD_LIBRARY = {
    "contract_title": {
        "label": "Sözleşme Başlığı",
        "type": "text",
        "placeholder": "Örn. Konut Kira Sözleşmesi",
        "example": "Konut Kira Sözleşmesi",
        "help": "Dokümanda üst başlık olarak kullanılır.",
        "required": False,
    },
    "contract_date": {
        "label": "Sözleşme Tarihi",
        "type": "date",
        "placeholder": "YYYY-AA-GG",
        "example": "2026-03-16",
        "help": "İmzaya esas tarih.",
        "required": True,
    },
    "effective_date": {
        "label": "Yürürlük Tarihi",
        "type": "date",
        "placeholder": "YYYY-AA-GG",
        "example": "2026-04-01",
        "help": "Sözleşmenin hüküm ifade etmeye başladığı tarih.",
        "required": False,
    },
    "party_a_name": {
        "label": "Taraf A Ünvan/Ad Soyad",
        "type": "text",
        "placeholder": "Örn. ABC Teknoloji A.Ş. / Ali Yılmaz",
        "example": "ABC Teknoloji A.Ş.",
        "help": "Şirket ise ünvan, kişi ise ad-soyad girin.",
        "required": True,
    },
    "party_a_type": {
        "label": "Taraf A Tipi",
        "type": "select",
        "options": ["Şirket", "Gerçek Kişi"],
        "placeholder": "",
        "example": "Şirket",
        "help": "Taraf A şirket mi gerçek kişi mi?",
        "required": True,
    },
    "party_a_tax_or_tc": {
        "label": "Taraf A Vergi/TC No",
        "type": "text",
        "placeholder": "Örn. 1234567890 / 12345678901",
        "example": "12345678901",
        "help": "Gerçek kişi TC (11 hane), şirket vergi no (10 hane).",
        "required": False,
    },
    "party_a_address": {
        "label": "Taraf A Adres",
        "type": "textarea",
        "placeholder": "Adres satırlarını girin",
        "example": "Maslak Mah. ... No: ... Şişli/İstanbul",
        "help": "Tebligat adresi olarak kullanılır.",
        "required": True,
    },
    "party_b_name": {
        "label": "Taraf B Ünvan/Ad Soyad",
        "type": "text",
        "placeholder": "Örn. XYZ Ltd. Şti. / Ayşe Demir",
        "example": "Ayşe Demir",
        "help": "Karşı taraf bilgisi.",
        "required": True,
    },
    "party_b_type": {
        "label": "Taraf B Tipi",
        "type": "select",
        "options": ["Şirket", "Gerçek Kişi"],
        "placeholder": "",
        "example": "Gerçek Kişi",
        "help": "Taraf B şirket mi gerçek kişi mi?",
        "required": True,
    },
    "party_b_tax_or_tc": {
        "label": "Taraf B Vergi/TC No",
        "type": "text",
        "placeholder": "Örn. 1234567890 / 12345678901",
        "example": "1234567890",
        "help": "Gerçek kişi TC (11 hane), şirket vergi no (10 hane).",
        "required": False,
    },
    "party_b_address": {
        "label": "Taraf B Adres",
        "type": "textarea",
        "placeholder": "Adres satırlarını girin",
        "example": "Çankaya Mah. ... No: ... Çankaya/Ankara",
        "help": "Tebligat adresi olarak kullanılır.",
        "required": True,
    },
    "currency": {
        "label": "Para Birimi",
        "type": "select",
        "options": ["TRY", "USD", "EUR"],
        "placeholder": "",
        "example": "TRY",
        "help": "Bedel/ödeme tutarları için para birimi.",
        "required": True,
    },
    "fee_amount": {
        "label": "Bedel / Ücret",
        "type": "number",
        "placeholder": "Örn. 50000",
        "example": "50000",
        "help": "KDV hariç/hariç değil bilgisi not alanında belirtilebilir.",
        "required": True,
    },
    "payment_terms": {
        "label": "Ödeme Şartı",
        "type": "textarea",
        "placeholder": "Örn. Aylık peşin, her ayın 5'ine kadar IBAN'a",
        "example": "Aylık peşin, her ayın 5'ine kadar",
        "help": "Vade, taksit, ödeme günü ve gecikme hükümleri.",
        "required": False,
    },
    "iban": {
        "label": "IBAN",
        "type": "text",
        "placeholder": "TR00 0000 0000 0000 0000 0000 00",
        "example": "TR12 0006 2000 1234 5678 9012 34",
        "help": "Ödeme yapılacak hesap.",
        "required": False,
    },
    "term_months": {
        "label": "Süre (Ay)",
        "type": "number",
        "placeholder": "Örn. 12",
        "example": "12",
        "help": "Belirli süreli sözleşmeler için.",
        "required": False,
    },
    "jurisdiction": {
        "label": "Yetkili Mahkeme",
        "type": "text",
        "placeholder": "Örn. İstanbul (Çağlayan) Mahkemeleri ve İcra Daireleri",
        "example": "İstanbul Mahkemeleri",
        "help": "Uyuşmazlık halinde yetkili yer.",
        "required": True,
    },
    "signature_city": {
        "label": "İmza Şehri",
        "type": "text",
        "placeholder": "Örn. İstanbul",
        "example": "İstanbul",
        "help": "İmza yeri.",
        "required": False,
    },
    "signature_count": {
        "label": "Nüsha Sayısı",
        "type": "number",
        "placeholder": "Örn. 2",
        "example": "2",
        "help": "Kaç nüsha hazırlanacağı.",
        "required": False,
    },
}


def _field(key: str, **overrides: Any) -> Dict[str, Any]:
    base = dict(FIELD_LIBRARY.get(key) or {})
    if not base:
        base = {"label": _humanize_key(key), "type": "text", "placeholder": "", "example": "", "help": "", "required": False}
    base["key"] = key
    base.update({k: v for k, v in overrides.items() if v is not None})
    return base


def _humanize_key(key: str) -> str:
    k = re.sub(r"[^a-zA-Z0-9_]+", " ", str(key or "")).strip().replace("_", " ")
    k = re.sub(r"\s+", " ", k).strip()
    return k[:1].upper() + k[1:] if k else "Alan"


def _infer_fields_from_content(content: str) -> List[Dict[str, Any]]:
    keys = []
    for m in re.finditer(r"\{\{\s*([^}]+?)\s*\}\}", content or ""):
        raw = (m.group(1) or "").strip()
        if not raw:
            continue
        key = raw.split("|", 1)[0].strip()
        if key and key not in keys:
            keys.append(key)
    out = []
    for k in keys:
        out.append(_field(k, label=_humanize_key(k), placeholder=f"Örn. { _humanize_key(k) }", help="Bu alan şablon içindeki ilgili bölümü doldurur."))
    return out


def _mk_contract_skeleton(title: str, purpose: str, fields: List[Dict[str, Any]], extra_clauses: Optional[List[str]] = None) -> str:
    extra = "\n".join([f"- {c}" for c in (extra_clauses or [])]).strip()
    extra = extra if extra else "-"
    field_lines = []
    for f in fields or []:
        k = str((f or {}).get("key") or "").strip()
        if not k:
            continue
        label = str((f or {}).get("label") or _humanize_key(k)).strip()
        field_lines.append(f"- {label}: {{{{{k}}}}}")
    info_block = "\n".join(field_lines).strip() if field_lines else "-"
    return (
        f"{title}\n\n"
        "TARİH: {{contract_date}}\n"
        "YÜRÜRLÜK: {{effective_date}}\n\n"
        "TARAFLAR:\n"
        "1) {{party_a_name}} ({{party_a_type}}) / Adres: {{party_a_address}} / Vergi/TC: {{party_a_tax_or_tc}}\n"
        "2) {{party_b_name}} ({{party_b_type}}) / Adres: {{party_b_address}} / Vergi/TC: {{party_b_tax_or_tc}}\n\n"
        f"KONU VE AMAÇ:\n{purpose}\n\n"
        f"DOLDURULACAK ALANLAR:\n{info_block}\n\n"
        "MADDELER:\n"
        "1. Tanımlar ve yorum\n"
        "2. Kapsam ve yükümlülükler\n"
        "3. Süre: {{term_months}} ay (varsa) / fesih koşulları\n"
        "4. Bedel/Ödeme: {{fee_amount}} {{currency}} / {{payment_terms}} / IBAN: {{iban}}\n"
        "5. Gizlilik, KVKK ve veri güvenliği\n"
        "6. Sorumluluk sınırlaması, mücbir sebep\n"
        "7. Uyuşmazlık çözümü ve yetki: {{jurisdiction}}\n\n"
        "TASLAK NOTLARI (opsiyonel ek hükümler):\n"
        f"{extra}\n\n"
        "İMZA:\n"
        "İmza Yeri: {{signature_city}}\n"
        "Nüsha: {{signature_count}}\n\n"
        "TARAF A: ______________________    TARAF B: ______________________\n"
    )


@lru_cache(maxsize=1)
def _seed_catalog() -> Dict[str, Any]:
    categories = list(SEED_CATEGORIES)

    templates: List[Dict[str, Any]] = []

    def add(cat_key: str, sub: str, title: str, description: str, extra_fields: List[Dict[str, Any]], extra_clauses: Optional[List[str]] = None):
        cat = next((c for c in categories if c["key"] == cat_key), None) or {}
        tid = f"{cat_key}_{re.sub(r'[^a-z0-9]+', '_', title.lower()).strip('_')}"
        base_fields = [
            _field("contract_date"),
            _field("effective_date", required=False),
            _field("party_a_name"),
            _field("party_a_type"),
            _field("party_a_address"),
            _field("party_a_tax_or_tc", required=False),
            _field("party_b_name"),
            _field("party_b_type"),
            _field("party_b_address"),
            _field("party_b_tax_or_tc", required=False),
            _field("term_months", required=False),
            _field("fee_amount"),
            _field("currency"),
            _field("payment_terms", required=False),
            _field("iban", required=False),
            _field("jurisdiction"),
            _field("signature_city", required=False),
            _field("signature_count", required=False),
        ]
        fields = base_fields + (extra_fields or [])
        content = _mk_contract_skeleton(title, description, fields, extra_clauses=extra_clauses)
        templates.append(
            {
                "id": tid,
                "title": title,
                "category": cat.get("title") or cat_key,
                "category_key": cat_key,
                "subcategory": sub,
                "category_icon": cat.get("icon"),
                "description": description,
                "content": content,
                "fields": fields,
            }
        )

    kira_fields = [
        _field("property_address", label="Kiralanan Taşınmaz Adresi", type="textarea", placeholder="Örn. ...", required=True),
        _field("rent_amount", label="Aylık Kira Bedeli", type="number", placeholder="Örn. 25000", required=True),
        _field("deposit_amount", label="Depozito", type="number", placeholder="Örn. 50000", required=False),
        _field("rent_payment_day", label="Ödeme Günü", type="number", placeholder="Örn. 5", required=False),
        _field("increase_method", label="Kira Artış Yöntemi", type="select", options=["TÜFE (12 aylık ort.)", "Sabit oran", "Taraf mutabakatı"], required=False),
        _field("furnished", label="Eşyalı mı?", type="checkbox", required=False),
    ]
    for t in [
        ("Konut", "Konut Kira Sözleşmesi", "TBK’ya uygun konut kiralaması için temel kira sözleşmesi.", ["Depozito iadesi, aidat, demirbaş listesi."]),
        ("İşyeri", "İşyeri Kira Sözleşmesi", "İşyeri kiralamalarında kullanım amacı, tadilat ve stopaj düzenlemeleri.", ["Stopaj ve yan giderler, tadilat onayı."]),
        ("Kısa Süre", "Kısa Süreli Kira Protokolü", "Kısa süreli kiralamalarda teslim, temizlik ve cayma şartları.", ["Teslim-tahliye tutanağı ekleri."]),
        ("Tahliye", "Tahliye Taahhütnamesi Protokolü", "Kiracının tahliye taahhüdü ve teslim koşulları.", ["Tahliye tarihi ve cezai şart."]),
        ("Depozito", "Depozito İade Protokolü", "Depozito iadesi/mahsup koşullarını düzenler.", ["Hasar tespiti ve mahsup kalemleri."]),
        ("Tadilat", "Tadilat ve Demirbaş Protokolü", "Tadilat izinleri, demirbaş listesi ve geri verme şartları.", ["Demirbaş listesi ve teslim fotoğrafları."]),
        ("Emlak", "Emlak Komisyon Sözleşmesi", "Emlak aracılık ve komisyon şartlarını düzenler.", ["Komisyon oranı ve doğum şartı."]),
        ("Ortak", "Ortak Alan Kullanım Protokolü", "Site/işyeri ortak alan kullanım kuralları ve sorumluluklar.", ["Hasar ve kullanım kısıtları."]),
    ]:
        add("kira", t[0], t[1], t[2], kira_fields, extra_clauses=t[3])

    is_fields = [
        _field("job_title", label="Pozisyon / Ünvan", type="text", required=True),
        _field("workplace_address", label="İşyeri Adresi", type="textarea", required=True),
        _field("salary_amount", label="Ücret", type="number", required=True),
        _field("salary_frequency", label="Ücret Periyodu", type="select", options=["Aylık", "Haftalık", "Günlük"], required=False),
        _field("start_date", label="İşe Başlama Tarihi", type="date", required=True),
        _field("working_hours", label="Çalışma Saatleri", type="text", required=False, placeholder="Örn. 09:00-18:00"),
        _field("probation_months", label="Deneme Süresi (Ay)", type="number", required=False),
        _field("remote_work", label="Uzaktan Çalışma", type="checkbox", required=False),
        _field("meal_allowance", label="Yemek Ücreti", type="number", required=False),
        _field("transport_allowance", label="Yol Ücreti", type="number", required=False),
    ]
    for t in [
        ("Belirsiz", "Belirsiz Süreli İş Sözleşmesi", "4857 sayılı İş Kanunu çerçevesinde belirsiz süreli iş sözleşmesi.", ["Fazla mesai, izin, disiplin süreçleri."]),
        ("Belirli", "Belirli Süreli İş Sözleşmesi", "Proje/işin niteliği gereği belirli süreli iş sözleşmesi.", ["Süre bitimi, uzatma koşulları."]),
        ("Kısmi", "Kısmi Süreli (Part-time) İş Sözleşmesi", "Kısmi süreli çalışma için ücret/çalışma saatleri düzenlemesi.", ["Haftalık saat ve puantaj."]),
        ("Uzaktan", "Uzaktan Çalışma Sözleşmesi", "Uzaktan çalışmada ekipman, veri güvenliği ve denetim hükümleri.", ["Ekipman zimmeti, KVKK."]),
        ("Deneme", "Deneme Süreli İş Protokolü", "Deneme süresi koşulları ve değerlendirme kriterleri.", ["Deneme sonunda bildirime ilişkin not."]),
        ("Gizlilik", "İşveren-Çalışan Gizlilik Protokolü", "İş ilişkisi kapsamında gizli bilgi ve sır saklama yükümlülüğü.", ["Cezai şart ve süre."]),
        ("Rekabet", "Rekabet Yasağı Sözleşmesi", "Çalışan için rekabet yasağı; süre/yer/konu sınırlarıyla.", ["Tazminat ve ölçülülük notu."]),
        ("Staj", "Stajyer Sözleşmesi", "Stajyer çalıştırılmasına ilişkin temel hükümler.", ["Eğitim planı ve sorumluluk."]),
        ("Taşeron", "Alt İşveren (Taşeron) Sözleşmesi", "Alt işveren ilişkisi ve sorumluluk paylaşımı düzenlemesi.", ["İSG, SGK, rücu."]),
        ("İbraname", "İşçi İbraname Protokolü", "İbraname ve ödeme kalemleri için düzenleme.", ["Ödeme dekontu ekleri."]),
    ]:
        add("is", t[0], t[1], t[2], is_fields, extra_clauses=t[3])

    hizmet_fields = [
        _field("service_scope", label="Hizmet Kapsamı", type="textarea", required=True),
        _field("service_location", label="Hizmet Yeri", type="text", required=False),
        _field("service_start", label="Başlangıç Tarihi", type="date", required=False),
        _field("service_end", label="Bitiş Tarihi", type="date", required=False),
        _field("sla_level", label="SLA Seviyesi", type="select", options=["Standart", "Gelişmiş", "Kritik"], required=False),
        _field("penalty_rate", label="Cezai Şart Oranı", type="number", required=False),
    ]
    for t in [
        ("Danışmanlık", "Danışmanlık Sözleşmesi", "Profesyonel danışmanlık hizmeti, kapsam, raporlama ve ücret düzenlemesi.", ["Teslimat formatı, toplantı sıklığı."]),
        ("Hizmet", "Hizmet Alım Sözleşmesi", "Belirli bir hizmetin yerine getirilmesi için kapsam ve ödeme şartları.", ["Kabul kriterleri ve teslim tutanağı."]),
        ("Bakım", "Bakım ve Destek Sözleşmesi", "Periyodik bakım, arıza müdahalesi ve destek süreçleri.", ["Müdahale süreleri, yedek parça."]),
        ("Temizlik", "Temizlik Hizmeti Sözleşmesi", "Temizlik personeli, vardiya ve iş güvenliği düzenlemeleri.", ["Malzeme temini, kontrol listesi."]),
        ("Güvenlik", "Özel Güvenlik Hizmeti Sözleşmesi", "Güvenlik hizmeti, devriye planı ve sorumluluklar.", ["Silah/ekipman, vardiya planı."]),
        ("Catering", "Yemek (Catering) Hizmeti Sözleşmesi", "Yemek hizmeti, menü, hijyen ve teslim koşulları.", ["Numune saklama, alerjen bildirimi."]),
        ("IT", "BT Yönetim Hizmeti Sözleşmesi", "BT operasyonları, bakım, güvenlik ve SLA düzenlemeleri.", ["Yetki matrisi, log kayıtları."]),
        ("Saha", "Saha Hizmetleri Sözleşmesi", "Saha kurulum, bakım, seyahat ve masraf düzenlemeleri.", ["Masraf kalemleri ve onay."]),
        ("Eğitim", "Kurumsal Eğitim Hizmeti Sözleşmesi", "Kurumsal eğitim programı, içerik ve eğitmen yükümlülükleri.", ["Katılımcı listesi, sertifika."]),
        ("Çağrı", "Çağrı Merkezi Hizmeti Sözleşmesi", "Çağrı merkezi hizmeti, kalite metrikleri ve veri güvenliği.", ["KVKK, çağrı kayıt saklama."]),
    ]:
        add("hizmet", t[0], t[1], t[2], hizmet_fields, extra_clauses=t[3])

    satis_fields = [
        _field("product_or_service", label="Ürün/Hizmet Tanımı", type="textarea", required=True),
        _field("quantity", label="Miktar", type="number", required=False),
        _field("delivery_date", label="Teslim Tarihi", type="date", required=False),
        _field("delivery_place", label="Teslim Yeri", type="text", required=False),
        _field("warranty_months", label="Garanti Süresi (Ay)", type="number", required=False),
        _field("incoterms", label="Teslim Şekli (Incoterms)", type="select", options=["EXW", "FCA", "CPT", "CIP", "DAP", "DDP"], required=False),
    ]
    for t in [
        ("Satış", "Mal Satış Sözleşmesi", "Mal satışı, teslim, ayıplı mal ve ödeme şartlarını düzenler.", ["KDV dahil/hariç, iade şartları."]),
        ("Hizmet", "Hizmet Satış Sözleşmesi", "Satılan hizmetin kapsamı, teslim ve kabul kriterleri.", ["Kabul testleri, revizyon sayısı."]),
        ("Tedarik", "Tedarik Sözleşmesi", "Süreklilik arz eden tedarik ilişkisi ve fiyat/teslim hükümleri.", ["Fiyat revizyonu ve stok planı."]),
        ("Bayi", "Bayi / Distribütörlük Sözleşmesi", "Bayi ağı, bölge, fiyat politikası ve rekabet hükümleri.", ["Hedefler ve prim sistemi."]),
        ("İade", "İade ve Değişim Protokolü", "İade/ürün değişimi süreçleri ve masraf paylaşımı.", ["Kargo sorumluluğu."]),
        ("Satış Sonrası", "Satış Sonrası Destek Protokolü", "Garanti, servis, yedek parça ve SLA koşulları.", ["Servis süreleri."]),
        ("Franchise", "Franchise Sözleşmesi", "Franchise kullanım hakkı, standartlar ve bedel düzenlemeleri.", ["Know-how, marka kullanımı."]),
        ("Kampanya", "Kampanya ve İndirim Protokolü", "Kampanya süresi, indirim oranı ve iade/iptal şartları.", ["Kampanya şartları metni."]),
        ("İhale", "Teklif ve Kabul Protokolü", "Teklif koşulları, kabul beyanı ve yürürlük hükümleri.", ["Teslim planı."]),
        ("Distribütör", "Yetkili Satıcı Sözleşmesi", "Yetkili satıcı şartları ve marka standartları.", ["Denetim ve eğitim."]),
    ]:
        add("satis", t[0], t[1], t[2], satis_fields, extra_clauses=t[3])

    gizlilik_fields = [
        _field("confidential_info_def", label="Gizli Bilgi Tanımı", type="textarea", required=False),
        _field("nda_term_months", label="Gizlilik Süresi (Ay)", type="number", required=False),
        _field("penalty_amount", label="Cezai Şart Tutarı", type="number", required=False),
        _field("ip_scope", label="Fikri Hak Kapsamı", type="textarea", required=False),
    ]
    for t in [
        ("NDA", "Gizlilik Sözleşmesi (NDA) - Karşılıklı", "Karşılıklı gizli bilgi paylaşımı için temel NDA.", ["İstisnalar, iade/imha, cezai şart."]),
        ("NDA", "Gizlilik Sözleşmesi (NDA) - Tek Taraflı", "Tek taraflı bilgi paylaşımı için NDA.", ["Paylaşım amacı ve kapsam."]),
        ("IP", "Fikri Hak Devri Sözleşmesi", "Üretilen eser/yazılım/tasarım haklarının devri için.", ["Devir bedeli ve kapsam."]),
        ("IP", "Lisans Sözleşmesi", "Fikri hakkın kullanım lisansı, süre, bölge ve alt lisans koşulları.", ["Lisans türü ve kısıtlar."]),
        ("KVKK", "Veri İşleme ve Gizlilik Ek Protokolü", "KVKK uyumu için veri işleme hükümleri.", ["Veri güvenliği teknik/idari tedbirler."]),
        ("Marka", "Marka Kullanım İzni Protokolü", "Marka kullanım koşulları ve denetim hükümleri.", ["Kılavuz ve onay süreçleri."]),
    ]:
        add("gizlilik", t[0], t[1], t[2], gizlilik_fields, extra_clauses=t[3])

    sirket_fields = [
        _field("company_name", label="Şirket Ünvanı", type="text", required=False),
        _field("share_percentage", label="Pay Oranı (%)", type="number", required=False),
        _field("share_price", label="Pay Bedeli", type="number", required=False),
        _field("board_powers", label="Yönetim Yetkileri", type="textarea", required=False),
        _field("non_compete_months", label="Rekabet Yasağı Süresi (Ay)", type="number", required=False),
    ]
    for t in [
        ("Ortaklık", "Ortaklık Sözleşmesi", "Ortaklar arası sermaye, yetki ve kâr paylaşımı düzenlemesi.", ["Kâr dağıtımı ve çıkış hükümleri."]),
        ("Pay Devri", "Pay Devri Sözleşmesi", "Pay devri, devir bedeli ve teminatlar.", ["Ayıba karşı tekeffül, kapanış şartları."]),
        ("Hissedar", "Hissedarlar Sözleşmesi", "Oy hakları, yönetim, veto konuları ve exit mekanizmaları.", ["Drag-along / tag-along."]),
        ("Yetki", "Yetki ve İmza Sirküleri Protokolü", "Temsil/ilzam yetkileri ve limitler.", ["Limit ve onay matrisi."]),
        ("Niyet", "Niyet Mektubu (LOI)", "Yatırım/satın alma görüşmeleri için bağlayıcılık sınırları.", ["Gizlilik ve münhasırlık."]),
        ("Münhasır", "Münhasırlık Protokolü", "Belirli süre münhasırlık ve ihlal yaptırımları.", ["Süre ve cezai şart."]),
        ("Yönetim", "Yönetim Hizmet Sözleşmesi", "Yönetim/operasyon hizmeti ve ücretlendirme düzenlemesi.", ["Performans göstergeleri."]),
        ("Kâr", "Kâr Paylaşım Protokolü", "Proje/iş birliği kâr paylaşımı ve mali raporlama.", ["Raporlama periyodu."]),
    ]:
        add("sirket", t[0], t[1], t[2], sirket_fields, extra_clauses=t[3])

    teknoloji_fields = [
        _field("project_scope", label="Proje Kapsamı", type="textarea", required=True),
        _field("deliverables", label="Teslimatlar", type="textarea", required=False),
        _field("acceptance_criteria", label="Kabul Kriterleri", type="textarea", required=False),
        _field("uptime", label="Uptime Taahhüdü (%)", type="number", required=False),
        _field("data_processing", label="Veri İşleme Kapsamı", type="textarea", required=False),
    ]
    for t in [
        ("SaaS", "SaaS Abonelik Sözleşmesi", "SaaS hizmeti, abonelik, kullanım limitleri ve SLA şartları.", ["Kullanım sınırları, aşım ücretleri."]),
        ("Geliştirme", "Yazılım Geliştirme Sözleşmesi", "Analiz-geliştirme-test teslim süreçleri ve kabul kriterleri.", ["Sprint planı, revizyon sayısı."]),
        ("SLA", "SLA (Hizmet Seviyesi) Ek Protokolü", "Uptime, müdahale süreleri, cezai şart ve raporlama.", ["Kritik/majör/minör tanımları."]),
        ("Bakım", "Yazılım Bakım ve Destek Sözleşmesi", "Sürüm güncelleme, güvenlik yamaları ve destek kanalları.", ["Destek saatleri."]),
        ("KVKK", "Veri İşleyen Sözleşmesi (DPA)", "KVKK/uyum için veri işleyen-veren yükümlülükleri.", ["Alt işlemciler ve denetim."]),
        ("Lisans", "Kurumsal Yazılım Lisans Sözleşmesi", "On-prem lisans, kullanım hakkı ve lisans denetimi.", ["Lisans sayısı, audit."]),
        ("API", "API Kullanım Koşulları Sözleşmesi", "API erişimi, rate limit, güvenlik ve sorumluluklar.", ["Anahtar güvenliği, IP kısıtları."]),
        ("Hosting", "Barındırma (Hosting) Sözleşmesi", "Sunucu barındırma, yedekleme, veri güvenliği ve SLA.", ["Yedekleme sıklığı."]),
    ]:
        add("teknoloji", t[0], t[1], t[2], teknoloji_fields, extra_clauses=t[3])

    finans_fields = [
        _field("principal_amount", label="Ana Para Tutarı", type="number", required=True),
        _field("interest_rate", label="Faiz Oranı (%)", type="number", required=False),
        _field("repayment_plan", label="Geri Ödeme Planı", type="textarea", required=False),
        _field("guarantor_name", label="Kefil Ad Soyad/Ünvan", type="text", required=False),
        _field("collateral_desc", label="Teminat Açıklaması", type="textarea", required=False),
    ]
    for t in [
        ("Ödünç", "Ödünç Para (Karz) Sözleşmesi", "Borç verme ilişkisi, vade ve geri ödeme planı.", ["Temerrüt ve masraflar."]),
        ("Kefalet", "Kefalet Sözleşmesi", "Kefil yükümlülüğü, limit ve süre hükümleri.", ["Müteselsil kefalet seçenekleri."]),
        ("Teminat", "Teminat Protokolü", "Teminatın türü, iade şartları ve paraya çevirme koşulları.", ["Rehin/temlik notları."]),
        ("Taksit", "Taksitlendirme Protokolü", "Borç için taksit planı ve temerrüt düzenlemesi.", ["Ödeme takvimi."]),
        ("Tahsilat", "Alacak Tahsilat Protokolü", "Tahsilat süreci, masraflar, mutabakat ve feragat hükümleri.", ["İcra takibi notu."]),
        ("Mutabakat", "Cari Hesap Mutabakat Protokolü", "Cari hesap, mutabakat ve itiraz süreleri.", ["Mutabakat periyodu."]),
        ("Temlik", "Alacak Temliki Sözleşmesi", "Alacağın devri, bildirim ve teminatlar.", ["Borçluya ihbar."]),
        ("Rehin", "Ticari İşletme Rehni Protokolü", "Rehin konusu, kayıt ve paraya çevirme koşulları.", ["Rehin tescili."]),
    ]:
        add("finans", t[0], t[1], t[2], finans_fields, extra_clauses=t[3])

    lojistik_fields = [
        _field("cargo_desc", label="Yük Tanımı", type="textarea", required=True),
        _field("pickup_place", label="Teslim Alma Yeri", type="text", required=False),
        _field("dropoff_place", label="Teslim Etme Yeri", type="text", required=False),
        _field("delivery_deadline", label="Teslim Süresi", type="text", required=False),
        _field("damage_liability", label="Hasar Sorumluluğu Notu", type="textarea", required=False),
    ]
    for t in [
        ("Taşıma", "Taşıma Sözleşmesi", "Yük taşımacılığı, teslim süreleri ve hasar sorumluluğu.", ["CMR/taşıma senedi notu."]),
        ("Depolama", "Depolama Sözleşmesi", "Depolama, giriş-çıkış ve sigorta hükümleri.", ["Stok mutabakatı."]),
        ("Dağıtım", "Dağıtım Hizmeti Sözleşmesi", "Bölgesel dağıtım, KPI ve hizmet seviyesi düzenlemesi.", ["KPI listesi."]),
        ("Kargo", "Kargo ve Teslim Protokolü", "Kargo masrafı, teslim koşulları ve iade prosedürü.", ["İade politikası."]),
        ("Lojistik", "3PL Lojistik Sözleşmesi", "3. parti lojistik hizmetleri, depo+taşıma birleşik.", ["Envanter sorumluluğu."]),
        ("Gümrük", "Gümrükleme Hizmeti Sözleşmesi", "Gümrük süreçleri ve belge sorumlulukları.", ["Vekaletname ve masraflar."]),
    ]:
        add("lojistik", t[0], t[1], t[2], lojistik_fields, extra_clauses=t[3])

    egitim_fields = [
        _field("program_name", label="Program / Hizmet Adı", type="text", required=True),
        _field("program_schedule", label="Takvim", type="textarea", required=False),
        _field("participant_count", label="Katılımcı Sayısı", type="number", required=False),
        _field("location", label="Yer", type="text", required=False),
        _field("cancellation_terms", label="İptal/İade Şartı", type="textarea", required=False),
    ]
    for t in [
        ("Kurs", "Kurs Sözleşmesi", "Kurs kaydı, ücret, devamsızlık ve sertifika şartları.", ["İptal-iade koşulları."]),
        ("Seminer", "Seminer / Eğitim Organizasyonu Sözleşmesi", "Eğitim organizasyonu, eğitmen ve içerik yükümlülükleri.", ["İçerik hakları."]),
        ("Etkinlik", "Etkinlik Organizasyon Sözleşmesi", "Etkinlik planı, mekan, tedarikçiler ve sorumluluklar.", ["Sigorta ve güvenlik."]),
        ("Sağlık", "Klinik Hizmet Sözleşmesi", "Sağlık hizmeti, randevu, aydınlatma ve ücret koşulları.", ["Aydınlatma metni."]),
        ("Atölye", "Atölye Çalışması Sözleşmesi", "Atölye içeriği, ekipman ve katılım şartları.", ["Malzeme listesi."]),
        ("Sponsorluk", "Sponsorluk Sözleşmesi", "Sponsorluk bedeli, görünürlük ve marka kullanımı.", ["Logo yerleşimleri."]),
    ]:
        add("egitim", t[0], t[1], t[2], egitim_fields, extra_clauses=t[3])

    pazarlama_fields = [
        _field("campaign_name", label="Kampanya Adı", type="text", required=True),
        _field("deliverable_count", label="İçerik Adedi", type="number", required=False),
        _field("platforms", label="Platformlar", type="textarea", required=False, placeholder="Örn. Instagram, TikTok"),
        _field("posting_dates", label="Paylaşım Tarihleri", type="textarea", required=False),
        _field("brand_guidelines", label="Marka Kılavuzu Notu", type="textarea", required=False),
    ]
    for t in [
        ("Influencer", "Influencer İş Birliği Sözleşmesi", "Influencer içerikleri, onay süreci ve ücret koşulları.", ["Etiket/hashtag zorunluluğu."]),
        ("Ajans", "Reklam Ajansı Hizmet Sözleşmesi", "Ajans hizmet kapsamı, teslimatlar ve raporlama.", ["Medya planı ekleri."]),
        ("Yayın", "Reklam Yayın Sözleşmesi", "Reklamın yayınlanması, süre, ölçümleme ve sorumluluklar.", ["KPI ve raporlama."]),
        ("İçerik", "İçerik Üretim Sözleşmesi", "Fotoğraf/video/metin üretimi, teslim ve hak devri/lisans.", ["Kullanım süresi ve mecralar."]),
        ("PR", "PR ve İletişim Danışmanlığı Sözleşmesi", "Basın iletişimi, kriz yönetimi ve raporlama.", ["Kriz planı."]),
        ("Medya", "Medya Satın Alma Sözleşmesi", "Medya satın alma, bütçe, komisyon ve raporlama.", ["Bütçe onay akışı."]),
        ("E-ticaret", "E-ticaret Satıcı Sözleşmesi", "Pazaryeri satıcı şartları, iade, komisyon ve KVKK hükümleri.", ["Komisyon oranları."]),
        ("Affiliate", "Affiliate / İş Ortaklığı Sözleşmesi", "Affiliate komisyonu, takip mekanizması ve ödeme şartları.", ["Komisyon tablosu."]),
        ("Abonelik", "Abonelik Koşulları Sözleşmesi", "Abonelik yenileme, iptal ve ücretlendirme şartları.", ["Otomatik yenileme bildirimi."]),
    ]:
        add("pazarlama", t[0], t[1], t[2], pazarlama_fields, extra_clauses=t[3])

    return {"categories": categories, "templates": templates}

# --- Schema ---

class ContractAnalysisRequest(BaseModel):
    title: str
    content: str

class TemplateCreate(BaseModel):
    title: str
    category: str
    content: str
    description: Optional[str] = None

class TemplateUpdate(BaseModel):
    title: str
    category: str
    content: str
    description: Optional[str] = None

class ContractGenerateRequest(BaseModel):
    template_id: str
    values: Dict[str, Any] = {}

class ContractCompareRequest(BaseModel):
    left_title: Optional[str] = None
    left_content: str
    right_title: Optional[str] = None
    right_content: str

class ContractClauseGenerateRequest(BaseModel):
    clause_type: str
    context: Optional[str] = None
    preferences: Dict[str, Any] = {}

class ContractExportRequest(BaseModel):
    title: str
    content: str
    format: str

class ContractAnalysisPreviewRequest(BaseModel):
    title: str
    content: str
    analysis: Dict[str, Any]
    company_name: Optional[str] = None

class ContractAnalysisExportRequest(ContractAnalysisPreviewRequest):
    format: str
    logo_url: Optional[str] = None


def _build_analysis_report_text(title: str, contract_text: str, analysis: Dict[str, Any], user: Dict[str, Any], company_name: str) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    uline = f"{(user.get('first_name') or user.get('firstName') or '').strip()} {(user.get('last_name') or user.get('lastName') or '').strip()}".strip()
    email = (user.get("email") or "").strip()
    who = uline or email or "-"
    company = (company_name or "").strip()
    header = []
    header.append(company if company else "Miron AI")
    header.append(f"Kullanıcı: {who}")
    if email and who != email:
        header.append(f"E-posta: {email}")
    header.append(f"Oluşturulma: {now}")
    header.append("")
    header.append(f"Rapor: {title}")
    header.append("=" * 48)

    def as_list(x):
        return x if isinstance(x, list) else []

    lines = []
    lines.extend(header)
    lines.append("")
    lines.append("GENEL ÖZET")
    lines.append("-" * 48)
    lines.append(str(analysis.get("genel_ozet") or "-"))
    rp = analysis.get("risk_puani")
    rs = analysis.get("risk_seviyesi")
    lines.append("")
    lines.append(f"RİSK PUANI: {rp if isinstance(rp, (int, float)) else '-'} / 100")
    lines.append(f"RİSK SEVİYESİ: {rs or '-'}")

    def section(title, items):
        lines.append("")
        lines.append(title)
        lines.append("-" * 48)
        it = as_list(items)
        if not it:
            lines.append("-")
        else:
            for i in it[:50]:
                lines.append(f"- {str(i)}")

    section("GÜÇLÜ YÖNLER", analysis.get("guclu_yonler"))
    section("ZAYIF/RİSKLİ YÖNLER", analysis.get("zayif_yonler"))
    section("GELECEK RİSKLERİ", analysis.get("gelecek_riskleri"))
    section("EKSİK MADDELER", analysis.get("eksik_maddeler"))
    section("ÖNERİLER", analysis.get("oneriler"))

    lines.append("")
    lines.append("HUKUKİ UYUM KONTROLLERİ")
    lines.append("-" * 48)
    uc = analysis.get("uyum_kontrolleri") if isinstance(analysis.get("uyum_kontrolleri"), dict) else {}
    if not uc:
        lines.append("-")
    else:
        for k, arr in uc.items():
            lines.append(f"{str(k).upper()}:")
            if not isinstance(arr, list) or not arr:
                lines.append("  -")
            else:
                for x in arr[:20]:
                    lines.append(f"  - {str(x)}")

    lines.append("")
    lines.append("MADDE BAZLI BULGULAR")
    lines.append("-" * 48)
    maddeler = analysis.get("maddeler") if isinstance(analysis.get("maddeler"), list) else []
    if not maddeler:
        lines.append("-")
    else:
        for m in maddeler[:40]:
            if not isinstance(m, dict):
                continue
            lines.append(f"* {m.get('baslik') or '-'} (risk: {m.get('risk') or '-'})")
            if m.get("gerekce"):
                lines.append(f"  Gerekçe: {m.get('gerekce')}")
            if m.get("onerilen_duzenleme"):
                lines.append(f"  Önerilen düzenleme: {m.get('onerilen_duzenleme')}")

    lines.append("")
    lines.append("ANALİZ EDİLEN SÖZLEŞME METNİ")
    lines.append("-" * 48)
    lines.append(contract_text.strip() or "-")
    return "\n".join(lines).strip() + "\n"


def _load_user_for_report(user_id: str) -> Dict[str, Any]:
    with get_db_cursor() as cur:
        cur.execute("SELECT id, email, first_name, last_name FROM users WHERE id = %s LIMIT 1", (user_id,))
        return dict(cur.fetchone() or {})


def _extract_text_from_upload(filename: str, data: bytes) -> str:
    fn = (filename or "").lower()

    if fn.endswith(".pdf"):
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                return "\n".join([(page.extract_text() or "") for page in pdf.pages])
        except Exception:
            return data.decode("utf-8", errors="ignore")

    if fn.endswith(".docx"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception:
            return data.decode("utf-8", errors="ignore")

    return data.decode("utf-8", errors="ignore")

# --- Endpoints ---

@router.get("/templates")
def list_templates(
    category: Optional[str] = None,
    include_remote: bool = Query(False),
    catalog: bool = Query(False),
    include_seed: bool = Query(True),
):
    """Mevcut sözleşme şablonlarını listele"""
    seed = _seed_catalog() if include_seed else {"categories": [], "templates": []}
    categories = list(seed.get("categories") or [])
    templates: List[Dict[str, Any]] = list(seed.get("templates") or [])

    sql = "SELECT * FROM contract_templates"
    params: tuple = ()
    if category and not catalog:
        sql += " WHERE category = %s"
        params = (category,)

    use_remote_db = (os.getenv("TEST_USE_REMOTE_DB", "false") or "").lower() == "true"
    if (os.getenv("ENVIRONMENT") or "").lower() != "test" or use_remote_db:
        with get_db_cursor() as cur:
            cur.execute(sql, params)
            rows = cur.fetchall() or []
            for r in rows:
                row = dict(r)
                row["id"] = str(row.get("id"))
                if "fields" not in row:
                    row["fields"] = _infer_fields_from_content(str(row.get("content") or ""))
                templates.append(row)

    if include_remote and (os.getenv("ENVIRONMENT") or "").lower() != "test":
        for r in _fetch_remote_templates():
            if not isinstance(r, dict):
                continue
            row = dict(r)
            row["id"] = str(row.get("id"))
            if "fields" not in row:
                row["fields"] = _infer_fields_from_content(str(row.get("content") or ""))
            templates.append(row)

    if category:
        c = str(category).strip().lower()
        templates = [
            t
            for t in templates
            if str(t.get("category_key") or "").lower() == c or str(t.get("category") or "").lower() == c
        ]

    if not catalog:
        out: List[Dict[str, Any]] = []
        for t in templates:
            if not isinstance(t, dict):
                continue
            row = dict(t)
            if "id" in row:
                row["id"] = str(row["id"])
            out.append(row)
        return out

    summaries: List[Dict[str, Any]] = []
    for t in templates:
        if not isinstance(t, dict):
            continue
        row = dict(t)
        fields = row.get("fields") or []
        summaries.append(
            {
                "id": str(row.get("id")),
                "title": row.get("title"),
                "category": row.get("category"),
                "category_key": row.get("category_key"),
                "category_icon": row.get("category_icon"),
                "subcategory": row.get("subcategory"),
                "description": row.get("description"),
                "fields": fields,
                "field_count": len(fields) if isinstance(fields, list) else 0,
            }
        )
    return {"categories": categories, "templates": summaries}


_REMOTE_CACHE = {"ts": 0.0, "items": []}

def _default_source_urls() -> List[str]:
    return [
        "https://raw.githubusercontent.com/mironintelligence/Miron22/main/backend/contract_templates_public.json",
    ]

def _fetch_remote_templates(category: Optional[str] = None) -> List[Dict[str, Any]]:
    ttl = int(os.getenv("CONTRACT_TEMPLATE_REMOTE_TTL_SECONDS", "900"))
    now = time.time()
    if _REMOTE_CACHE["items"] and (now - float(_REMOTE_CACHE["ts"])) < ttl:
        items = list(_REMOTE_CACHE["items"])
        if category:
            items = [t for t in items if str(t.get("category") or "").lower() == str(category).lower()]
        return items

    raw = (os.getenv("CONTRACT_TEMPLATE_SOURCE_URLS") or "").strip()
    urls = [u.strip() for u in raw.split(",") if u.strip()] if raw else _default_source_urls()
    gathered: List[Dict[str, Any]] = []
    for u in urls:
        try:
            req = Request(u, headers={"User-Agent": "miron-ai"})
            with urlopen(req, timeout=8) as resp:
                data = resp.read().decode("utf-8", errors="replace")
            parsed = json.loads(data)
            if isinstance(parsed, list):
                for item in parsed:
                    if isinstance(item, dict):
                        gathered.append(item)
        except (URLError, HTTPError, ValueError):
            continue

    _REMOTE_CACHE["ts"] = now
    _REMOTE_CACHE["items"] = gathered

    if category:
        gathered = [t for t in gathered if str(t.get("category") or "").lower() == str(category).lower()]
    return gathered

@router.get("/templates/{template_id}")
def get_template(template_id: str):
    """Tekil şablon detayı"""
    tid = str(template_id)
    # DB check
    if tid.isdigit():
        with get_db_cursor() as cur:
            cur.execute("SELECT * FROM contract_templates WHERE id = %s", (int(tid),))
            row = cur.fetchone()
            if row:
                r = dict(row)
                r["id"] = str(r.get("id"))
                r.setdefault("fields", _infer_fields_from_content(str(r.get("content") or "")))
                return r

    for t in (_seed_catalog().get("templates") or []):
        if str(t.get("id")) == tid:
            row = dict(t)
            row["id"] = str(row.get("id"))
            return row

    for t in MOCK_TEMPLATES:
        if str(t.get("id")) == tid:
            row = dict(t)
            row["id"] = str(row.get("id"))
            row.setdefault("fields", _infer_fields_from_content(str(row.get("content") or "")))
            return row

    for t in _fetch_remote_templates():
        if str(t.get("id")) == tid:
            row = dict(t)
            row["id"] = str(row.get("id"))
            row.setdefault("fields", _infer_fields_from_content(str(row.get("content") or "")))
            return row
    
    raise HTTPException(status_code=404, detail="Şablon bulunamadı.")


@router.post("/generate")
def generate_contract(payload: ContractGenerateRequest, user: Dict[str, Any] = Depends(get_current_user)):
    client = get_openai_client()
    if not client:
        raise HTTPException(status_code=500, detail="AI servisi kullanılamıyor.")

    tpl = get_template(payload.template_id)
    content = str(tpl.get("content") or "")
    title = str(tpl.get("title") or "Sözleşme")
    values = payload.values or {}

    prompt = f"""
Sen Türk hukukuna hakim, titiz bir sözleşme yazım uzmanısın.
Aşağıdaki şablonu kullanarak, verilen alan değerleri ile eksiksiz bir sözleşme metni üret.
Şablondaki {{...}} alanlarını uygun şekilde doldur.
Çelişen veya eksik bilgi varsa, kullanıcıya kısa bir not bırak ve makul varsayım yapma.

SÖZLEŞME BAŞLIĞI: {title}

ALAN DEĞERLERİ (JSON):
{json.dumps(values, ensure_ascii=False)}

ŞABLON:
{content[:15000]}
"""

    try:
        completion = chat_completions_create(client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Çıktı sadece sözleşme metni olsun. Markdown kod bloğu kullanma."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        generated = (completion.choices[0].message.content or "").strip()

        sql = """
            INSERT INTO user_contracts (user_id, title, generated_content, created_at)
            VALUES (%s, %s, %s, NOW())
            RETURNING id
        """
        with get_db_cursor() as cur:
            cur.execute(sql, (user["id"], title, generated))
            row = cur.fetchone()
            cid = row["id"] if row else None

        return {"id": cid, "generated": generated}
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=500, detail="Sözleşme oluşturulamadı.")

@router.post("/analyze")
def analyze_contract(payload: ContractAnalysisRequest, user: Dict[str, Any] = Depends(get_current_user)):
    """
    AI ile Sözleşme Analizi (Risk, Güçlü/Zayıf Yönler)
    """
    client = get_openai_client()
    if not client:
        return {"error": "AI servisi kullanılamıyor."}

    prompt = f"""
Sen Türk hukukuna hakim, titiz bir sözleşme analistisin. Aşağıdaki sözleşme metnini detaylı analiz et.

SÖZLEŞME BAŞLIĞI: {payload.title}

METİN:
{payload.content[:15000]}

ÇIKTI KURALLARI:
- Sadece JSON üret.
- Belirsizlik varsa varsayım yapma; "eksik_bilgi" alanında belirt.
- Doğrudan metinden alıntı yaparken kısa tut.

İSTENEN JSON ŞEMASI:
{{
  "genel_ozet": "string",
  "risk_puani": 0,
  "risk_seviyesi": "dusuk|orta|yuksek",
  "guclu_yonler": ["string"],
  "zayif_yonler": ["string"],
  "gelecek_riskleri": ["string"],
  "eksik_maddeler": ["string"],
  "uyum_kontrolleri": {{
    "tbk": ["string"],
    "kvkk": ["string"],
    "ticaret": ["string"],
    "is_hukuku": ["string"]
  }},
  "maddeler": [
    {{
      "baslik": "string",
      "risk": "dusuk|orta|yuksek",
      "gerekce": "string",
      "onerilen_duzenleme": "string"
    }}
  ],
  "oneriler": ["string"],
  "eksik_bilgi": ["string"]
}}
"""
    
    try:
        completion = chat_completions_create(client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Sen Türk hukukuna hakim, titiz bir sözleşme analistisin."},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        analysis = json.loads(completion.choices[0].message.content)
        
        # Sonucu DB'ye kaydet (Opsiyonel - Kullanıcı geçmişi için)
        sql = """
            INSERT INTO user_contracts (user_id, title, original_content, analysis_result)
            VALUES (%s, %s, %s, %s)
            RETURNING id
        """
        with get_db_cursor() as cur:
             cur.execute(sql, (user['id'], payload.title, payload.content, json.dumps(analysis)))
             contract_id = cur.fetchone()['id']
             
        return {"id": contract_id, "title": payload.title, "content": payload.content, "analysis": analysis}

    except Exception as e:
        raise HTTPException(status_code=500, detail="Analiz yapılamadı.")


@router.post("/analyze-file")
async def analyze_contract_file(
    file: UploadFile = File(...),
    title: str = Form("Sözleşme Analizi"),
    user: Dict[str, Any] = Depends(get_current_user),
):
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Dosya çok büyük.")

    text = _extract_text_from_upload(file.filename or "", data)
    text = (text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Dosyadan metin çıkarılamadı.")
    return analyze_contract(ContractAnalysisRequest(title=title or "Sözleşme Analizi", content=text), user)


@router.post("/compare")
def compare_contracts(payload: ContractCompareRequest, user: Dict[str, Any] = Depends(get_current_user)):
    client = get_openai_client()
    if not client:
        raise HTTPException(status_code=500, detail="AI servisi kullanılamıyor.")

    left_title = payload.left_title or "Sürüm A"
    right_title = payload.right_title or "Sürüm B"
    prompt = f"""
Sen Türk hukukuna hakim bir sözleşme revizyon uzmanısın.
İki sözleşme sürümünü karşılaştır; değişiklikleri özetle, risk etkisini değerlendir.

SÜRÜM A BAŞLIĞI: {left_title}
SÜRÜM B BAŞLIĞI: {right_title}

SÜRÜM A:
{payload.left_content[:12000]}

SÜRÜM B:
{payload.right_content[:12000]}

Sadece JSON üret:
{{
  "ozet": "string",
  "degisiklikler": [
    {{
      "baslik": "string",
      "degisiklik": "string",
      "risk_etkisi": "azaldi|degismedi|artti",
      "onerilen_aksiyon": "string"
    }}
  ],
  "risk_karsilastirma": {{
    "a_risk_puani": 0,
    "b_risk_puani": 0,
    "degisim": "azaldi|degismedi|artti"
  }},
  "notlar": ["string"]
}}
"""
    try:
        completion = chat_completions_create(client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Sadece JSON döndür."},
                {"role": "user", "content": prompt},
            ],
            response_format={"type": "json_object"},
            temperature=0.2,
        )
        data = json.loads(completion.choices[0].message.content)
        return {"compare": data}
    except Exception:
        raise HTTPException(status_code=500, detail="Karşılaştırma yapılamadı.")


@router.post("/clauses/generate")
def generate_clause(payload: ContractClauseGenerateRequest, user: Dict[str, Any] = Depends(get_current_user)):
    client = get_openai_client()
    if not client:
        raise HTTPException(status_code=500, detail="AI servisi kullanılamıyor.")

    clause_type = (payload.clause_type or "").strip()
    if not clause_type:
        raise HTTPException(status_code=400, detail="clause_type gerekli.")
    ctx = (payload.context or "").strip()
    prefs = payload.preferences or {}

    prompt = f"""
Sen Türk hukukuna uygun sözleşme maddesi yazan bir uzmansın.
İstenen madde tipine göre bir sözleşme maddesi yaz.
Metin açık, uygulanabilir ve riskleri azaltacak şekilde olsun.

MADDE TİPİ: {clause_type}

TERCİHLER (JSON):
{json.dumps(prefs, ensure_ascii=False)}

BAĞLAM (varsa sözleşme taslağı / ilgili bölüm):
{ctx[:8000]}
"""
    try:
        completion = chat_completions_create(client,
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Çıktı sadece madde metni olsun. Markdown kod bloğu kullanma."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )
        clause = (completion.choices[0].message.content or "").strip()
        return {"clause": clause}
    except Exception:
        raise HTTPException(status_code=500, detail="Madde oluşturulamadı.")


@router.post("/analysis/preview")
def preview_analysis_report(payload: ContractAnalysisPreviewRequest, user: Dict[str, Any] = Depends(get_current_user)):
    company = (payload.company_name or os.getenv("DEFAULT_COMPANY_NAME") or "").strip()
    u = _load_user_for_report(str(user.get("id")))
    report = _build_analysis_report_text(payload.title, payload.content, payload.analysis, u or user, company)
    return {"report": report}


@router.post("/analysis/export")
def export_analysis_report(payload: ContractAnalysisExportRequest, user: Dict[str, Any] = Depends(get_current_user)):
    fmt = (payload.format or "").strip().lower()
    if fmt not in {"txt", "docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Geçersiz format.")

    company = (payload.company_name or os.getenv("DEFAULT_COMPANY_NAME") or "").strip()
    logo_url = (payload.logo_url or os.getenv("DEFAULT_COMPANY_LOGO_URL") or "").strip()
    u = _load_user_for_report(str(user.get("id")))
    report = _build_analysis_report_text(payload.title, payload.content, payload.analysis, u or user, company)

    import hashlib
    from fastapi.responses import Response
    from uyap_udf import _safe_filename

    safe = _safe_filename(payload.title or "analiz_raporu")
    ts = datetime.now().strftime("%Y%m%d%H%M")

    logo_bytes = b""
    if logo_url:
        try:
            req = Request(logo_url, headers={"User-Agent": "miron-ai"})
            with urlopen(req, timeout=8) as resp:
                logo_bytes = resp.read()
        except Exception:
            logo_bytes = b""

    if fmt == "txt":
        data = report.encode("utf-8")
        sha = hashlib.sha256(data).hexdigest()
        return Response(
            content=data,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.txt"', "X-File-SHA256": sha},
        )

    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import mm
            from reportlab.lib.utils import ImageReader
        except Exception:
            raise HTTPException(status_code=501, detail="PDF export devre dışı.")

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        x = 15 * mm
        y = h - 20 * mm

        if logo_bytes:
            try:
                img = ImageReader(io.BytesIO(logo_bytes))
                c.drawImage(img, x, y - 20 * mm, width=30 * mm, height=18 * mm, mask="auto")
                y -= 24 * mm
            except Exception:
                pass

        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, (company or "Miron AI")[:80])
        y -= 8 * mm
        c.setFont("Helvetica", 9)
        for line in report.splitlines():
            if y < 20 * mm:
                c.showPage()
                c.setFont("Helvetica", 9)
                y = h - 20 * mm
            c.drawString(x, y, line[:160])
            y -= 4.5 * mm
        c.save()

        data = buf.getvalue()
        sha = hashlib.sha256(data).hexdigest()
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.pdf"', "X-File-SHA256": sha},
        )

    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except Exception:
        raise HTTPException(status_code=501, detail="DOCX export devre dışı.")

    doc = Document()
    if logo_bytes:
        try:
            doc.add_picture(io.BytesIO(logo_bytes), width=Inches(1.2))
        except Exception:
            pass
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_paragraph(company or "Miron AI").runs[0].bold = True
    doc.add_paragraph(payload.title or "Analiz Raporu").runs[0].bold = True
    for line in report.splitlines():
        doc.add_paragraph(line)
    stream = io.BytesIO()
    doc.save(stream)
    data = stream.getvalue()
    sha = hashlib.sha256(data).hexdigest()
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.docx"', "X-File-SHA256": sha},
    )


@router.post("/export")
def export_contract(payload: ContractExportRequest, user: Dict[str, Any] = Depends(get_current_user)):
    fmt = (payload.format or "").strip().lower()
    title = (payload.title or "sozlesme").strip()[:200]
    content = (payload.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="İçerik boş.")
    if fmt not in {"docx", "pdf", "udf", "uyap"}:
        raise HTTPException(status_code=400, detail="Geçersiz format.")

    import io
    import hashlib
    from fastapi.responses import Response
    from uyap_udf import _build_udf, _safe_filename

    safe = _safe_filename(title or "sozlesme")
    ts = datetime.now().strftime("%Y%m%d%H%M")

    if fmt == "udf":
        udf_text = _build_udf(title, content, None)
        data = udf_text.encode("utf-8")
        sha = hashlib.sha256(data).hexdigest()
        return Response(
            content=data,
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.udf"', "X-File-SHA256": sha},
        )

    if fmt == "uyap":
        data = content.encode("utf-8")
        sha = hashlib.sha256(data).hexdigest()
        return Response(
            content=data,
            media_type="application/octet-stream",
            headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.uyap"', "X-File-SHA256": sha},
        )

    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import mm
        except Exception:
            raise HTTPException(status_code=501, detail="PDF export devre dışı.")

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        x = 15 * mm
        y = h - 20 * mm
        c.setFont("Helvetica", 10)
        for line in content.splitlines():
            if y < 20 * mm:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = h - 20 * mm
            c.drawString(x, y, line[:150])
            y -= 5 * mm
        c.save()
        data = buf.getvalue()
        sha = hashlib.sha256(data).hexdigest()
        return Response(
            content=data,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe}_{ts}.pdf"', "X-File-SHA256": sha},
        )

    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        raise HTTPException(status_code=501, detail="DOCX export devre dışı.")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    doc.add_paragraph(title).runs[0].bold = True
    for para in content.split("\n\n"):
        p = doc.add_paragraph(para.strip())
        p.paragraph_format.space_after = Pt(6)
    stream = io.BytesIO()
    doc.save(stream)
    data = stream.getvalue()
    sha = hashlib.sha256(data).hexdigest()
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename=\"{safe}_{ts}.docx\"', "X-File-SHA256": sha},
    )

@router.post("/templates", dependencies=[Depends(require_admin)])
def create_template(payload: TemplateCreate):
    """Admin: Yeni şablon ekle"""
    sql = """
        INSERT INTO contract_templates (title, category, content, description)
        VALUES (%s, %s, %s, %s)
        RETURNING id
    """
    with get_db_cursor() as cur:
        cur.execute(sql, (payload.title, payload.category, payload.content, payload.description))
        return {"id": cur.fetchone()['id'], "message": "Şablon oluşturuldu."}


@router.put("/templates/{template_id}", dependencies=[Depends(require_admin)])
def update_template(template_id: str, payload: TemplateUpdate):
    """Admin: Sözleşme şablonunu güncelle"""
    tid = str(template_id).strip()
    if not tid.isdigit():
        raise HTTPException(status_code=400, detail="Geçersiz template id.")

    sql = """
        UPDATE contract_templates
        SET title = %s,
            category = %s,
            content = %s,
            description = %s
        WHERE id = %s
        RETURNING id
    """
    with get_db_cursor() as cur:
        cur.execute(
            sql,
            (
                payload.title,
                payload.category,
                payload.content,
                payload.description,
                int(tid),
            ),
        )
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Şablon bulunamadı.")
        return {"ok": True, "id": str(row.get("id")), "message": "Şablon güncellendi."}


@router.delete("/templates/{template_id}", dependencies=[Depends(require_admin)])
def delete_template(template_id: str):
    """Admin: Sözleşme şablonunu sil"""
    tid = str(template_id).strip()
    if not tid.isdigit():
        raise HTTPException(status_code=400, detail="Geçersiz template id.")

    sql = """
        DELETE FROM contract_templates
        WHERE id = %s
        RETURNING id
    """
    with get_db_cursor() as cur:
        cur.execute(sql, (int(tid),))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Şablon bulunamadı.")
        return {"ok": True, "id": str(row.get("id")), "message": "Şablon silindi."}
