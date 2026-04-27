from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Query, Depends
from user_auth import get_current_user
from typing import Optional, List, Dict, Any
from datetime import datetime
import asyncio
import os, io, json

try:
    from openai_client import get_openai_client
    from services.risk_engine import risk_engine
    from security import sanitize_text
    from llm_gateway import chat_completions_create
except ImportError:
    from openai_client import get_openai_client
    from services.risk_engine import risk_engine
    from security import sanitize_text
    from llm_gateway import chat_completions_create

import pdfplumber
from docx import Document

router = APIRouter(prefix="/api/risk", tags=["Risk & Strateji Analizi"])

# Use advanced model for simulation
SIMULATION_MODEL = "gpt-4o"

# Upload hardening. Keep in sync with the frontend validator in Risk.jsx.
_ALLOWED_RISK_EXTS = {".pdf", ".docx", ".txt"}
_MAX_RISK_BYTES = 15 * 1024 * 1024  # 15 MB


async def _read_upload_limited(file: UploadFile) -> bytes:
    """Read an upload while enforcing the size cap.

    Without this, a malicious client can stream arbitrary bytes into server
    memory and trigger OOM. We also validate the extension against a strict
    allowlist — the LLM pipeline does not need to parse anything else.
    """
    name = (file.filename or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Dosya adı gerekli.")
    import os as _os
    ext = _os.path.splitext(name)[1].lower()
    if ext not in _ALLOWED_RISK_EXTS:
        raise HTTPException(
            status_code=415,
            detail="Yalnızca PDF, DOCX veya TXT desteklenir.",
        )
    buf = bytearray()
    chunk_size = 64 * 1024
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        if len(buf) + len(chunk) > _MAX_RISK_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Dosya çok büyük. Maksimum {_MAX_RISK_BYTES // (1024 * 1024)} MB.",
            )
        buf.extend(chunk)
    return bytes(buf)

@router.post("/simulate", response_model=dict)
async def simulate_case(
    case_description: Optional[str] = Form(None),
    jurisdiction: str = Form("Türkiye"),
    user_role: str = Form("Davacı"),
    file: Optional[UploadFile] = File(None),
    _user: dict = Depends(get_current_user),
):
    """
    Advanced Case Simulation with Deep Reasoning.
    """
    client = get_openai_client()
    if not client:
        raise HTTPException(status_code=500, detail="AI Client init failed.")
    
    # 1. Dosya varsa oku ve metne ekle (size-limited, type-restricted).
    file_meta: Dict[str, Any] = {
        "file_attached": False,
        "file_name": None,
        "file_text_chars_extracted": 0,
        "file_text_chars_embedded": 0,
    }
    file_content = ""
    if file:
        raw = await _read_upload_limited(file)
        safe_name = _safe_basename(file.filename or "")
        extracted = (extract_text_from_bytes(safe_name, raw) or "").strip()
        file_meta["file_attached"] = True
        file_meta["file_name"] = safe_name or None
        file_meta["file_text_chars_extracted"] = len(extracted)
        if not extracted:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Yüklenen dosyadan okunabilir metin çıkarılamadı. "
                    "PDF şifreli veya taranmış görüntü olabilir; DOCX/TXT deneyin."
                ),
            )
        embedded = extracted[:12000]
        file_meta["file_text_chars_embedded"] = len(embedded)
        file_content = (
            f"\n\n[EK DOSYA — simülasyonda kullanılacak: {safe_name}]\n{embedded}"
        )

    text_part = (case_description or "").strip()
    if not text_part and not file_content:
        raise HTTPException(status_code=400, detail="Dava metni veya dosya gereklidir.")

    full_text = (text_part + file_content).strip()
    clean_case = sanitize_text(full_text, 15000)
    
    det_risk = risk_engine.analyze_risk(clean_case)
    det_score = det_risk.get("risk_score", 50)
    det_issues = det_risk.get("key_issues", [])

    prompt = f"""
    Sen kıdemli bir stratejik dava danışmanısın.
    Aşağıdaki dava senaryosunu derinlemesine simüle et.
    "[EK DOSYA" ile başlayan bölüm varsa, bu metni senaryonun ayrılmaz parçası say;
    içindeki somut iddia, tarih, tutar ve şartları analizinde açıkça dikkate al.
    
    SENARYO:
    {clean_case}
    
    Taraf: {user_role}
    Yargı Yeri: {jurisdiction}
    
    DETERMİNİSTİK RİSK ANALİZİ BULGULARI (Referans Al):
    - Hesaplanan Temel Risk Skoru: {det_score}/100
    - Tespit Edilen Kritik Hususlar: {', '.join(det_issues)}
    
    GÖREVLERİN:
    1. Yargı Yeri ve Usul Analizi: Hangi mahkeme görevli? İspat yükü kimde?
    2. Risk Analizi: Zayıf noktalarımız neler? Karşı taraf ne diyebilir? (Deterministik bulguları yorumla)
    3. Simülasyon:
       - En İyi Senaryo: (Kazanma ihtimali, süre, maliyet)
       - En Kötü Senaryo: (Kaybetme riski, masraflar)
       - En Olası Sonuç: (Gerekçeli tahmin)
    4. Stratejik Tavsiye: Taktik, savunma, karşı-strateji ve uzlaşma planı üret.
    
    ZORUNLU YAPISAL KATMANLAR (JSON içinde mutlaka yer almalı):
    - procedural_risk: {{ "level": "Yüksek/Orta/Düşük", "details": "..." }}
    - contradiction_analysis: {{ "internal": "...", "external": "..." }}
    - missing_claims: [ "iddia1", "iddia2" ]
    - alternative_qualification: {{ "current": "...", "proposed": "...", "advantage": "..." }}
    
    ÇIKTI FORMATI (JSON):
    {{
        "jurisdiction_analysis": "...",
        "burden_of_proof": "...",
        "risk_factors": ["risk1", "risk2"],
        "counter_arguments": ["arg1", "arg2"],
        "scenarios": {{
            "best_case": "...",
            "worst_case": "...",
            "most_probable": "..."
        }},
        "procedural_risk": {{ "level": "...", "details": "..." }},
        "contradiction_analysis": {{ "internal": "...", "external": "..." }},
        "missing_claims": ["..."],
        "alternative_qualification": {{ "current": "...", "proposed": "...", "advantage": "..." }},
        "tactical_strategy": ["..."],
        "defensive_strategy": ["..."],
        "counter_strategy": ["..."],
        "settlement_analysis": ["..."],
        "probability_logic": "...",
        "win_probability_percent": 60,
        "estimated_duration_months": 12,
        "strategic_recommendation": "..."
    }}
    Sadece JSON döndür.
    """
    
    try:
        # Blocking OpenAI call inside an async endpoint — offload to a thread
        # so the event loop remains free for other concurrent requests.
        completion = await asyncio.to_thread(
            chat_completions_create,
            client,
            model=SIMULATION_MODEL,
            messages=[
                {"role": "system", "content": "Kıdemli bir stratejik hukuk danışmanısın. Sadece geçerli JSON üret. Uydurma yapma. Tüm metinler Türkçe olmalı."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            response_format={"type": "json_object"},
        )
        result = json.loads(completion.choices[0].message.content)
        file_meta["prompt_chars_after_sanitize"] = len(clean_case or "")
        file_meta["description_chars"] = len(text_part)
        result["simulation_input_meta"] = file_meta
        return result
    except Exception:
        # Never leak the raw exception chain to the client in production,
        # and avoid the unused-`e` pyflakes warning.
        raise HTTPException(status_code=500, detail="Simulation failed") from None

# ------- Helpers -------

def _safe_basename(name: str) -> str:
    return os.path.basename(name or "").strip()

def extract_text_from_bytes(filename: str, b: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        try:
            with pdfplumber.open(io.BytesIO(b)) as pdf:
                return "\n".join([p.extract_text() or "" for p in pdf.pages])
        except Exception:
            return b.decode("utf-8", errors="ignore")
    elif name.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(b))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join((cell.text or "").strip() for cell in row.cells))
            return "\n".join(parts)
        except Exception:
            return b.decode("utf-8", errors="ignore")
    return b.decode("utf-8", errors="ignore")

def guess_case_type(text: str) -> str:
    t = (text or "").lower()
    if any(k in t for k in ["boşan", "nafaka", "velayet"]):
        return "Aile (Boşanma)"
    if any(k in t for k in ["tazminat", "alacak", "haksız fiil"]):
        return "Tazminat"
    if any(k in t for k in ["iş mahkemesi", "işe iade", "kıdem", "ihbar"]):
        return "İş Hukuku"
    if any(k in t for k in ["icra", "ödeme emri", "takip dosya"]):
        return "İcra-İflas"
    if any(k in t for k in ["şikayet", "cumhuriyet başsavcılığı", "tck"]):
        return "Ceza"
    return "Genel"

def analyze_risk(text: str) -> Dict[str, Any]:
    return risk_engine.analyze_risk(text)

# ------- Endpoints -------

@router.post("/analyze")
async def risk_analyze(
    file: Optional[UploadFile] = File(None),
    case_text: Optional[str] = Form(None),
    first_name: Optional[str] = Form(None),
    last_name: Optional[str] = Form(None),
    _user: dict = Depends(get_current_user),
):
    """
    Yüklenen dosya veya case_text üzerinden risk puanı ve strateji önerisi üretir.
    """
    if not file and not (case_text and case_text.strip()):
        raise HTTPException(status_code=400, detail="Dosya veya metin (case_text) gereklidir.")

    if file:
        raw = await _read_upload_limited(file)
        safe_name = _safe_basename(file.filename or "")
        text = extract_text_from_bytes(safe_name, raw)
        source = safe_name or "dosya"
    else:
        text = case_text.strip()
        source = "metin"

    text = sanitize_text(text, 12000)
    result = analyze_risk(text)
    result.update({
        "source": source,
        "case_type_guess": guess_case_type(text),
        "length": len(text),
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    })

    return result
