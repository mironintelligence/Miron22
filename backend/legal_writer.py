from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, Any, Optional, Literal, List

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from uyap_udf import _build_udf, _safe_filename

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openai_client import get_openai_client

writer_router = APIRouter(prefix="/writer", tags=["Dilekçe Oluşturucu"])

# client None olmasın diye try/except:
# key yoksa preview/export patlar ama catalog/fields çalışır.
try:
    client = get_openai_client()
except Exception:
    client = None


# -----------------------------
# FIELD BUILDERS (DRY)
# -----------------------------
def f_text(key: str, label: str, required: bool = True, placeholder: str = ""):
    obj = {"key": key, "label": label, "type": "text", "required": required}
    if placeholder:
        obj["placeholder"] = placeholder
    return obj

def f_area(key: str, label: str, required: bool = True, placeholder: str = ""):
    obj = {"key": key, "label": label, "type": "textarea", "required": required}
    if placeholder:
        obj["placeholder"] = placeholder
    return obj

def base_court_party_fields() -> List[Dict[str, Any]]:
    return [
        f_text("court_name", "Mahkeme / Makam", True, "İstanbul Anadolu 7. Asliye Hukuk / ..."),
        f_text("city", "Şehir", True, "İstanbul"),
        f_text("plaintiff", "Başvuran / Davacı / Talep Eden", True),
        f_text("defendant", "Karşı Taraf / Davalı / Muhatap", False),
        f_text("subject", "Konu", True, "Alacak / Tespit / Tazminat / ..."),
        f_area("facts", "Vakıalar (Özet)", True),
        f_area("evidence", "Deliller", False, "Sözleşme, fatura, ekran görüntüsü, tanık, bilirkişi..."),
        f_area("requests", "Sonuç ve İstem", True, "Asıl/fer’i talepler, yargılama giderleri, vekalet ücreti..."),
    ]

def base_prosecutor_fields() -> List[Dict[str, Any]]:
    return [
        f_text("prosecutor", "Cumhuriyet Başsavcılığı", True, "İstanbul C.Başsavcılığı"),
        f_text("complainant", "Şikâyetçi", True),
        f_text("suspect", "Şüpheli (Varsa)", False),
        f_area("facts", "Olayın Özeti", True),
        f_area("offences", "İsnat Edilen Suçlar", True, "TCK m...."),
        f_area("evidence", "Deliller", False),
        f_area("requests", "Talep", True, "Soruşturma yapılıp kamu davası açılması..."),
    ]


# -----------------------------
# 60+ TEMPLATE CATALOG
# -----------------------------
def tpl(key: str, title: str, case_type: str, policy: str, fields: List[Dict[str, Any]]):
    return {"key": key, "title": title, "case_type": case_type, "policy": policy, "fields": fields}

CATALOG: Dict[str, Any] = {
    # 1
    "Adli Dilekçeler": [
        tpl("dava-genel", "Dava Dilekçesi (Genel)", "Medeni",
            "HMK usul; taraflar, konu, vakıalar, hukuki sebepler, deliller, sonuç-istem zorunlu.",
            base_court_party_fields()),
        tpl("cevap-dilekcesi", "Cevap / Savunma Dilekçesi", "Medeni",
            "HMK m.129; itiraz/defi, deliller ve sonuç-istem düzeni.",
            base_court_party_fields()),
        tpl("cevaba-cevap", "Cevaba Cevap Dilekçesi", "Medeni",
            "HMK süre ve usul; yeni vakıa/delil sınırları.",
            base_court_party_fields()),
        tpl("ikinci-cevap", "İkinci Cevap Dilekçesi", "Medeni",
            "HMK; son beyan düzeni, delil değerlendirme.",
            base_court_party_fields()),
        tpl("feragat", "Feragat Dilekçesi", "Medeni",
            "Feragat kesin sonuç doğurur; HMK usulüne uygun açık irade gerekir.",
            base_court_party_fields()),
        tpl("islah", "Islah Dilekçesi", "Medeni",
            "HMK ıslah koşulları; talep artırım/ değişiklik sınırları.",
            base_court_party_fields()),
        tpl("ihtiyati-tedbir", "İhtiyati Tedbir Talebi", "Medeni",
            "HMK m.389 vd; telafisi güç zarar, haklı görünüm, teminat.",
            base_court_party_fields()),
        tpl("delil-tespiti", "Delil Tespiti Talebi", "Medeni",
            "HMK delil tespiti; kaybolma/elde edilmesi zor deliller.",
            base_court_party_fields()),
        tpl("tanik-listesi", "Tanık Listesi Sunumu", "Medeni",
            "Tanık isim/adres ve ispat konusu açıkça yazılır.",
            base_court_party_fields()),
        tpl("bilirkisi-ek-soru", "Bilirkişi Raporuna Ek Soru / İtiraz", "Medeni",
            "Bilirkişi raporuna somut itiraz ve ek inceleme talebi.",
            base_court_party_fields()),
    ],

    # 2
    "Kanun Yolları": [
        tpl("istinaf", "İstinaf Dilekçesi", "Medeni",
            "HMK m.341 vd; hukuka aykırılık başlıkları, kaldırma/değiştirme istemi.",
            base_court_party_fields() + [f_text("decision_no", "Karar No/Tarih", True, "2025/.. E. 2025/.. K. / ...")]),
        tpl("temyiz", "Temyiz Dilekçesi", "Medeni",
            "HMK m.361 vd; bozma sebepleri açık ve maddeli olmalı.",
            base_court_party_fields() + [f_text("bam_decision", "BAM Karar No/Tarih", True)]),
        tpl("yargilamanin-iadesi", "Yargılamanın İadesi Talebi", "Medeni",
            "HMK; iade sebepleri ve süre şartları.",
            base_court_party_fields()),
        tpl("karar-duzeltme", "Karar Düzeltme Talebi", "Medeni",
            "Yargı yolu şartları ve sınırlı sebepler.",
            base_court_party_fields()),
        tpl("itiraz", "İtiraz Dilekçesi (Genel)", "Genel",
            "İtiraz sebepleri somut ve delille destekli olmalı.",
            base_court_party_fields()),
    ],

    # 3
    "Ceza Hukuku": [
        tpl("suc-duyurusu", "Suç Duyurusu Dilekçesi", "Ceza",
            "CMK usul; olay, suç vasfı, deliller, talep.",
            base_prosecutor_fields()),
        tpl("tutuklamaya-itiraz", "Tutuklamaya İtiraz", "Ceza",
            "CMK; ölçülülük, somut gerekçe yokluğu, adli kontrol yeterliliği.",
            [
                f_text("court_name", "Sulh Ceza Hakimliği", True),
                f_text("suspect", "Şüpheli/Sanık", True),
                f_text("file_no", "Soruşturma/Dosya No", True),
                f_area("grounds", "İtiraz Gerekçeleri", True),
                f_area("requests", "Talep", True),
            ]),
        tpl("adli-kontrol-itiraz", "Adli Kontrol Kararına İtiraz", "Ceza",
            "CMK; yükümlülüklerin ölçüsüzlüğü, değişiklik talebi.",
            [
                f_text("court_name", "Mahkeme/Hakimlik", True),
                f_text("suspect", "Şüpheli/Sanık", True),
                f_text("file_no", "Dosya No", True),
                f_area("grounds", "Gerekçeler", True),
                f_area("requests", "Talep", True),
            ]),
        tpl("hagb-itiraz", "HAGB Kararına İtiraz", "Ceza",
            "HAGB şartları ve itiraz gerekçeleri.",
            base_court_party_fields()),
        tpl("siber-suc-duyurusu", "Siber Suç Duyurusu", "Ceza",
            "Bilişim delilleri, log/IP/ekran görüntüsü ile desteklenmeli.",
            base_prosecutor_fields()),
    ],

    # 4
    "İcra-İflas": [
        tpl("icra-takibi-baslatma", "İcra Takibi Başlatma Talebi", "İcra",
            "İİK; alacak kalemi, faiz, dayanak belge.",
            [
                f_text("enforcement_office", "İcra Dairesi", True),
                f_text("creditor", "Alacaklı", True),
                f_text("debtor", "Borçlu", True),
                f_text("claim", "Alacak Miktarı ve Türü", True, "TL/Döviz, faiz..."),
                f_area("basis", "Dayanak", False, "Senet, fatura, sözleşme..."),
                f_area("requests", "Talep", True),
            ]),
        tpl("icra-itiraz", "İcra Takibine İtiraz", "İcra",
            "İİK m.62 vd; süresinde borca/imzaya/yetkiye itiraz.",
            [
                f_text("enforcement_office", "İcra Dairesi", True),
                f_text("debtor", "Borçlu", True),
                f_text("file_no", "Takip No", True),
                f_area("grounds", "İtiraz Gerekçeleri", True),
                f_area("requests", "Talep", True),
            ]),
        tpl("itirazin-iptali", "İtirazın İptali Davası", "İcra",
            "Alacağın ispatı, icra inkar tazminatı koşulları.",
            base_court_party_fields()),
        tpl("menfi-tespit", "Menfi Tespit Davası", "İcra",
            "Borçlu olmadığının tespiti; tedbir talebi değerlendirilebilir.",
            base_court_party_fields()),
        tpl("istihkak", "İstihkak İddiası Dilekçesi", "İcra",
            "Haczedilen mala üçüncü kişi mülkiyet iddiası.",
            base_court_party_fields()),
        tpl("iflas-erteleme", "İflas / Konkordato Ön Talep (Taslak)", "İcra",
            "Genel taslak: mali tablo, alacaklılar, ödeme planı istenir.",
            base_court_party_fields()),
    ],

    # 5
    "Aile Hukuku": [
        tpl("bosanma", "Boşanma Dilekçesi (Genel)", "Aile",
            "TMK; kusur, nafaka, tazminat, velayet, deliller.",
            base_court_party_fields() + [
                f_text("marriage_date", "Evlilik Tarihi", False),
                f_text("children", "Çocuklar (Varsa)", False),
            ]),
        tpl("nafaka-artirim", "Nafaka Artırım Davası", "Aile",
            "Gelir değişimi, ihtiyaç artışı, hakkaniyet.",
            base_court_party_fields()),
        tpl("velayet-degisikligi", "Velayet Değişikliği Davası", "Aile",
            "Çocuğun üstün yararı, yaşam koşulları, deliller.",
            base_court_party_fields()),
        tpl("koruma-karari", "6284 Koruma Kararı Talebi", "Aile",
            "Şiddet/tehdit iddiaları, acil tedbir.",
            base_court_party_fields()),
        tpl("mal-rejimi", "Mal Rejimi Tasfiyesi Dilekçesi", "Aile",
            "Edinilmiş mallar/katkı payı/alacak kalemleri.",
            base_court_party_fields()),
    ],

    # 6
    "İş Hukuku": [
        tpl("ise-iade", "İşe İade Dava Dilekçesi", "İş",
            "4857; fesih geçersizliği, süre, işe iade/boşta geçen süre.",
            base_court_party_fields() + [
                f_text("employment_period", "Çalışma Süresi", False),
                f_text("last_wage", "Son Ücret", False),
            ]),
        tpl("kidem-ihbar", "Kıdem / İhbar Tazminatı Talebi", "İş",
            "Kıdem/ihbar, fazla mesai, yıllık izin alacakları.",
            base_court_party_fields()),
        tpl("fazla-mesai", "Fazla Mesai Alacağı Talebi", "İş",
            "Puantaj/kayıtlar, tanık, bordro itirazı.",
            base_court_party_fields()),
        tpl("mobbing", "Mobbing (Psikolojik Taciz) Dava Dilekçesi", "İş",
            "Süreç, deliller, tazminat kalemleri.",
            base_court_party_fields()),
        tpl("ise-baslatmama", "İşe Başlatmama Tazminatı Talebi", "İş",
            "İşe iade sonrası işe başlatmama sonuçları.",
            base_court_party_fields()),
    ],

    # 7
    "Tüketici Hukuku": [
        tpl("ayipli-mal", "Ayıplı Mal İadesi / Bedel İadesi", "Tüketici",
            "6502; ayıp, seçimlik haklar, servis raporu/fatura delili.",
            base_court_party_fields()),
        tpl("ayipli-hizmet", "Ayıplı Hizmet Şikayeti / Dava", "Tüketici",
            "Hizmetin ayıplı ifası, bedel iadesi/tazminat.",
            base_court_party_fields()),
        tpl("mesafeli-satis", "Mesafeli Satış Cayma Bildirimi", "Tüketici",
            "Cayma hakkı, iade süreci, yazılı bildirim.",
            base_court_party_fields()),
        tpl("tuketici-hakem", "Tüketici Hakem Heyeti Başvurusu", "Tüketici",
            "Başvuru özeti, ekler, talep.",
            base_court_party_fields()),
        tpl("kredi-kart-itiraz", "Kredi Kartı Harcama İtirazı", "Tüketici",
            "Yetkisiz işlem/itiraz, chargeback delilleri.",
            base_court_party_fields()),
    ],

    # 8
    "İdare Hukuku": [
        tpl("idari-para-cezasi-itiraz", "İdari Para Cezasına İtiraz", "İdare",
            "Usul ve maddi hata, deliller, iptal talebi.",
            base_court_party_fields()),
        tpl("vergi-ceza-uzlasma", "Vergi Ceza İhbarnamesi Uzlaşma Talebi", "İdare",
            "Uzlaşma gerekçesi, ödeme planı yaklaşımı.",
            base_court_party_fields()),
        tpl("iptal-davasi", "İptal Davası Dilekçesi", "İdare",
            "İdari işlemin hukuka aykırılığı; yürütmenin durdurulması talebi.",
            base_court_party_fields()),
        tpl("tam-yargi", "Tam Yargı Davası Dilekçesi", "İdare",
            "Zarar, illiyet, hizmet kusuru/ kusursuz sorumluluk.",
            base_court_party_fields()),
        tpl("bilgi-edinme", "Bilgi Edinme Başvurusu", "İdare",
            "4982 kapsamında bilgi talebi.",
            base_court_party_fields()),
    ],

    # 9
    "Gayrimenkul & Tapu": [
        tpl("tapu-iptal-tescil", "Tapu İptal ve Tescil Davası", "Gayrimenkul",
            "Muris muvazaası/ehliyetsizlik vs. somut olay.",
            base_court_party_fields()),
        tpl("ecrimisil", "Ecrimisil (Haksız İşgal Tazminatı)", "Gayrimenkul",
            "Haksız kullanım süresi, emsal kira, deliller.",
            base_court_party_fields()),
        tpl("ortakligin-giderilmesi", "Ortaklığın Giderilmesi (İzale-i Şuyu)", "Gayrimenkul",
            "Paydaşlık, satış/aynı taksim koşulları.",
            base_court_party_fields()),
        tpl("tahliye", "Tahliye Dava Dilekçesi (Genel)", "Gayrimenkul",
            "Kira sözleşmesi, tahliye sebebi, ihtar/delil.",
            base_court_party_fields()),
        tpl("kira-tespit", "Kira Tespit Davası", "Gayrimenkul",
            "Emsal kira, rayiç, bilirkişi.",
            base_court_party_fields()),
    ],

    # 10
    "Ticaret & Şirketler": [
        tpl("alacak-ticari", "Ticari Alacak Davası", "Ticaret",
            "Fatura/cari hesap/sözleşme delilleri.",
            base_court_party_fields()),
        tpl("ihtarname-odememe", "Ödeme İhtarname Taslağı", "Ticaret",
            "Borç ve süre verilmesi, temerrüt.",
            base_court_party_fields()),
        tpl("haksiz-rekabet", "Haksız Rekabet Davası", "Ticaret",
            "TTK; eylem, zarar, tespit/men/tazminat.",
            base_court_party_fields()),
        tpl("marka-ihlali", "Marka İhlali İhtar / Dava Taslağı", "Ticaret",
            "Marka hakkı, ihlal örnekleri, delil.",
            base_court_party_fields()),
        tpl("sirket-genel-kurul-iptal", "Genel Kurul Karar İptali", "Ticaret",
            "TTK; usulsüz çağrı/oy/kapsam.",
            base_court_party_fields()),
    ],

    # 11
    "KVKK & Bilişim": [
        tpl("kvkk-basvuru", "KVKK Veri Sorumlusuna Başvuru", "KVKK",
            "KVKK m.11 talepleri; silme/düzeltme/erişim.",
            base_court_party_fields()),
        tpl("kvkk-sikayet", "KVKK Kurul Şikayeti (Taslak)", "KVKK",
            "Başvuru yapıldıysa ekle; ihlal ve deliller.",
            base_court_party_fields()),
        tpl("erisim-engeli", "Erişim Engeli / İçerik Kaldırma Talebi", "Bilişim",
            "5651; URL, ihlal içeriği, hak ihlali.",
            base_court_party_fields() + [f_text("url", "URL/Link", True)]),
        tpl("hakaret-siber", "Hakaret / Tehdit (Dijital Delil) Dilekçesi", "Ceza",
            "Ekran görüntüsü, link, tarih-saat delilleri.",
            base_prosecutor_fields()),
        tpl("telif-ihlal", "Telif Hakkı İhlali İhtar Taslağı", "Bilişim",
            "Eser sahipliği, ihlal örnekleri, kaldırma talebi.",
            base_court_party_fields()),
    ],

    # 12
    "Arabuluculuk": [
        tpl("arabuluculuk-basvuru", "Arabuluculuk Başvuru Dilekçesi", "Arabuluculuk",
            "Uyuşmazlık özeti, taraf bilgileri, talep.",
            base_court_party_fields()),
        tpl("arabuluculuk-itiraz", "Arabuluculuk Son Tutanağına İlişkin Not (Taslak)", "Arabuluculuk",
            "Süreç notu/ek beyan taslağı.",
            base_court_party_fields()),
        tpl("is-arabuluculuk", "İş Uyuşmazlığı Arabuluculuk Başvurusu", "Arabuluculuk",
            "İşçilik alacakları özeti ve talepler.",
            base_court_party_fields()),
        tpl("ticari-arabuluculuk", "Ticari Arabuluculuk Başvurusu", "Arabuluculuk",
            "Cari hesap/fatura/temerrüt özeti.",
            base_court_party_fields()),
        tpl("tuketici-arabuluculuk", "Tüketici Arabuluculuk Başvurusu", "Arabuluculuk",
            "Ayıplı mal/hizmet uyuşmazlığı özeti.",
            base_court_party_fields()),
    ],

    # 13 (ek kategori, 60+ garanti)
    "Noter & İhtar": [
        tpl("noter-ihtar", "Noter İhtarname Taslağı (Genel)", "Noter",
            "Muhataba süre verme, temerrüt, kayıt altına alma.",
            base_court_party_fields()),
        tpl("teslim-tutanak", "Teslim / İade Tutanağı Taslağı", "Noter",
            "Teslim edilen mal/hizmet, tarih-saat, imza alanı.",
            base_court_party_fields()),
        tpl("ihtar-kira", "Kira İhtarname (Ödeme / Tahliye)", "Noter",
            "Kira borcu, süre, tahliye ihtarı.",
            base_court_party_fields()),
        tpl("ihtar-hizmet", "Hizmet Sözleşmesi Fesih Bildirimi", "Noter",
            "Fesih nedeni, tarih, alacak/borç.",
            base_court_party_fields()),
        tpl("vekalet-azil", "Vekalet Azli Bildirimi", "Noter",
            "Azil iradesi ve bildirim.",
            base_court_party_fields()),
    ],
}


# ---------------- Models ---------------- #
class PreviewRequest(BaseModel):
    template_key: str
    values: Dict[str, Any]
    language: Literal["TR", "EN"] = "TR"
    include_statutes: bool = True
    include_case_law: bool = True
    mask_pii: bool = True

class ExportRequest(PreviewRequest):
    format: Literal["docx", "uyap", "udf"] = "docx"
# ---------------- Helpers ---------------- #
def get_template_by_key(key: str) -> Optional[Dict[str, Any]]:
    for cat_items in CATALOG.values():
        for t in cat_items:
            if t["key"] == key:
                return t
    return None

def build_prompt(tpl_obj: Dict[str, Any], req: PreviewRequest) -> str:
    lang = "Turkish" if req.language == "TR" else "English"

    flags = []
    if req.include_statutes:
        flags.append("ADD_STATUTES")
    if req.include_case_law:
        flags.append("ADD_CASELAW")
    if req.mask_pii:
        flags.append("MASK_PII")
    flags_str = ",".join(flags) if flags else "NONE"

    def kvs(d: Dict[str, Any]) -> str:
        out = []
        for k, v in (d or {}).items():
            if v is None or str(v).strip() == "":
                continue
            out.append(f"- {k}: {v}")
        return "\n".join(out) if out else "-"

    return f"""
You are Libra AI Legal Writer. Produce formal legal petitions (dilekçe) conforming to Turkish legal style.
Output must include these sections and ONLY these markers in order:
#HEADER
#SUMMARY
#LEGAL_BASIS
#RESULT_REQUESTS
#ATTACHMENTS

#LANGUAGE: {lang}
#CASE_TYPE: {tpl_obj["case_type"]}
#POLICY: {tpl_obj["policy"]}
#FLAGS: {flags_str}

#FORM_VALUES
{kvs(req.values)}

#INSTRUCTIONS
- Use formal tone; typical Turkish legal petition structure.
- If ADD_STATUTES, cite relevant articles (HMK/TMK/TBK/İİK/CMK/4857/6502/5651/KVKK) where appropriate.
- If ADD_CASELAW, mention 1-2 concise Yargıtay precedent hints (no long quotes).
- RESULT_REQUESTS must include attorney fee + costs when appropriate.
""".strip()

def parse_marked_text(text: str) -> Dict[str, str]:
    keys = ["#HEADER", "#SUMMARY", "#LEGAL_BASIS", "#RESULT_REQUESTS", "#ATTACHMENTS"]
    out = {"header": "", "summary": "", "legal_basis": "", "result_requests": "", "attachments": ""}
    s = (text or "").replace("\r\n", "\n")

    idx = {k: s.find(k) for k in keys if s.find(k) != -1}
    if not idx:
        out["summary"] = s.strip()
        return out

    ordered = sorted(idx.items(), key=lambda x: x[1])
    for i, (mk, start) in enumerate(ordered):
        end = ordered[i + 1][1] if i + 1 < len(ordered) else len(s)
        body = s[start + len(mk): end].strip()
        if mk == "#HEADER":
            out["header"] = body
        elif mk == "#SUMMARY":
            out["summary"] = body
        elif mk == "#LEGAL_BASIS":
            out["legal_basis"] = body
        elif mk == "#RESULT_REQUESTS":
            out["result_requests"] = body
        elif mk == "#ATTACHMENTS":
            out["attachments"] = body
    return out


# ---------------- Routes ---------------- #
@writer_router.get("/catalog")
def get_catalog():
    # Sadece liste görünümü: fields yok
    data = []
    for cat, items in CATALOG.items():
        data.append({
            "category": cat,
            "items": [{"key": t["key"], "title": t["title"], "case_type": t["case_type"]} for t in items]
        })
    return data

@writer_router.get("/fields/{template_key}")
def get_fields(template_key: str):
    tpl_obj = get_template_by_key(template_key)
    if not tpl_obj:
        raise HTTPException(status_code=404, detail="Şablon bulunamadı")
    return {"key": tpl_obj["key"], "title": tpl_obj["title"], "fields": tpl_obj["fields"]}

@writer_router.post("/preview")
def preview(req: PreviewRequest):
    tpl_obj = get_template_by_key(req.template_key)
    if not tpl_obj:
        raise HTTPException(status_code=400, detail="Geçersiz template_key")

    if client is None:
        raise HTTPException(status_code=500, detail="OpenAI client yok. OPENAI_API_KEY kontrol et ve restart at.")

    prompt = build_prompt(tpl_obj, req)
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Libra AI Legal Writer"},
                {"role": "user", "content": prompt},
            ],
        )
        draft = resp.choices[0].message.content or ""
        return parse_marked_text(draft)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM/Preview: {e}")
@writer_router.post("/export")
def export_doc(req: ExportRequest):
    tpl_obj = get_template_by_key(req.template_key)
    if not tpl_obj:
        raise HTTPException(status_code=400, detail="Geçersiz template_key")

    if client is None:
        raise HTTPException(status_code=500, detail="OpenAI client yok. OPENAI_API_KEY kontrol et ve restart at.")

    fmt = (req.format or "docx").strip().lower()

    prompt = build_prompt(tpl_obj, req)
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Libra AI Legal Writer"},
                {"role": "user", "content": prompt},
            ],
        )
        final_text = resp.choices[0].message.content or ""
        sections = parse_marked_text(final_text)

        # ✅ UYAP / UDF: düz metin indir
        if fmt in ("uyap", "udf"):
            text_out = "\n\n".join([
                sections.get("header", "").strip(),
                sections.get("summary", "").strip(),
                "HUKUKİ SEBEPLER:\n" + sections.get("legal_basis", "").strip(),
                "SONUÇ VE İSTEM:\n" + sections.get("result_requests", "").strip(),
                "EKLER:\n" + sections.get("attachments", "").strip(),
            ]).strip()

            stream = io.BytesIO(text_out.encode("utf-8"))
            filename = f"dilekce_{tpl_obj['key']}{datetime.now().strftime('%Y%m%d%H%M')}.{fmt}"

            return StreamingResponse(
                stream,
                media_type="application/octet-stream",
                headers={"Content-Disposition": f'attachment; filename="{filename}"'},
            )

        # ✅ DOCX: aynı eski mantık
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title_p = doc.add_paragraph(tpl_obj["title"])
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title_p.runs:
            title_p.runs[0].bold = True
            title_p.runs[0].font.size = Pt(14)

        def add_block(h: str, body: str):
            ph = doc.add_paragraph(h)
            run = ph.runs[0] if ph.runs else ph.add_run(h)
            run.bold = True
            for para in (body or "").split("\n\n"):
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(6)

        add_block("Başlık", sections.get("header", ""))
        add_block("Özet", sections.get("summary", ""))
        add_block("Hukuki Sebepler", sections.get("legal_basis", ""))
        add_block("Sonuç ve İstem", sections.get("result_requests", ""))
        add_block("Ekler", sections.get("attachments", ""))

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        filename = f"dilekce_{tpl_obj['key']}{datetime.now().strftime('%Y%m%d%H%M')}.docx"
        return StreamingResponse(
            stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export: {e}")